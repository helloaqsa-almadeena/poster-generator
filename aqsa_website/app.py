from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import pandas as pd

app = Flask(__name__)

# ── Constants ──────────────────────────────────────────────────────────────
SHELF_ROW_H  = 1872
SHELF_COL0_W = 11779
SHELF_COL1_W = 3075
SHELF_TBL_W  = SHELF_COL0_W + SHELF_COL1_W

# ── XML helpers ────────────────────────────────────────────────────────────
def make_run(text, font_ascii, sz_halfpt, bold=True, color_rgb=None):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_ascii)
    rFonts.set(qn('w:hAnsi'), font_ascii)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if color_rgb:
        ce = OxmlElement('w:color')
        ce.set(qn('w:val'), f'{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}')
        rPr.append(ce)
    sz = OxmlElement('w:sz');   sz.set(qn('w:val'), str(sz_halfpt));  rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(sz_halfpt)); rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r

def set_para_center_no_space(p):
    pPr = p._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:after'), '0'); sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)

def page_break(doc):
    p = doc.add_paragraph()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)

def split_price(price):
    integer = int(float(price))
    decimal = f"{float(price):.2f}".split(".")[1]
    return integer, decimal

# ── A4 Portrait ────────────────────────────────────────────────────────────
def create_a4_portrait_single(items, currency="د.إ"):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0); s.page_height = Cm(29.7)
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(1.27)

    for idx, item in enumerate(items):
        name = item["name"].upper()
        integer, decimal = split_price(item["price"])

        # Blank spacer
        spacer = doc.add_paragraph()
        pPr = spacer._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
        sp.set(qn('w:line'), '6480'); sp.set(qn('w:lineRule'), 'exact')
        pPr.append(sp)
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'both'); pPr.append(jc)

        # Price line: integer + .decimal + د.إ
        p_price = doc.add_paragraph()
        set_para_center_no_space(p_price)
        p_price._p.append(make_run(str(integer),  'Impact', 500, bold=True, color_rgb=(0xEE, 0x00, 0x00)))
        p_price._p.append(make_run(f'.{decimal}', 'Impact', 130, bold=True, color_rgb=(0xEE, 0x00, 0x00)))
        p_price._p.append(make_run(f' {currency}','aed',     96, bold=True, color_rgb=(0x00, 0x00, 0x00)))

        # Product name
        p_name = doc.add_paragraph()
        set_para_center_no_space(p_name)
        p_name._p.append(make_run(name, 'Impact', 72, bold=True, color_rgb=(0x00, 0x00, 0x00)))

        if idx < len(items) - 1:
            page_break(doc)

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ── Shelf label helpers ────────────────────────────────────────────────────
def _shelf_border_xml():
    borders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'10')
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),'000000')
        borders.append(el)
    return borders

def _shelf_cell_margins():
    mar = OxmlElement('w:tcMar')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:w'),'100'); el.set(qn('w:type'),'dxa')
        mar.append(el)
    return mar

def _add_shelf_row(table, name, integer, decimal, currency):
    row = table.add_row(); tr = row._tr
    trPr = OxmlElement('w:trPr')
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:hRule'),'exact'); trH.set(qn('w:val'), str(SHELF_ROW_H))
    trPr.append(trH); tr.insert(0, trPr)

    c0 = row.cells[0]
    tcP0 = c0._tc.get_or_add_tcPr()
    w0 = OxmlElement('w:tcW'); w0.set(qn('w:w'), str(SHELF_COL0_W)); w0.set(qn('w:type'),'dxa')
    tcP0.insert(0, w0); tcP0.append(_shelf_border_xml()); tcP0.append(_shelf_cell_margins())
    va0 = OxmlElement('w:vAlign'); va0.set(qn('w:val'),'center'); tcP0.append(va0)
    p0 = c0.paragraphs[0]; set_para_center_no_space(p0)
    p0._p.append(make_run(name, 'Britannic Bold', 72, bold=True, color_rgb=(0,0,0)))

    c1 = row.cells[1]
    tcP1 = c1._tc.get_or_add_tcPr()
    w1 = OxmlElement('w:tcW'); w1.set(qn('w:w'), str(SHELF_COL1_W)); w1.set(qn('w:type'),'dxa')
    tcP1.insert(0, w1); tcP1.append(_shelf_border_xml()); tcP1.append(_shelf_cell_margins())
    va1 = OxmlElement('w:vAlign'); va1.set(qn('w:val'),'center'); tcP1.append(va1)
    p1 = c1.paragraphs[0]; set_para_center_no_space(p1)
    p1._p.append(make_run(str(integer),  'Britannic Bold', 190, bold=True, color_rgb=(0xC8,0x00,0x00)))
    p1._p.append(make_run(f'.{decimal}', 'Britannic Bold',  72, bold=True, color_rgb=(0xC8,0x00,0x00)))
    p1._p.append(make_run(f' {currency}','aed',             60, bold=True, color_rgb=(0xC8,0x00,0x00)))

def _new_shelf_table(doc):
    table = doc.add_table(rows=0, cols=2)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(SHELF_TBL_W)); tblW.set(qn('w:type'),'dxa')
    tblPr.insert(0, tblW)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
    return table

def create_a5_landscape_shelf(items, currency="د.إ"):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0); s.page_height = Cm(14.8)
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(1.0)
    sectPr = s._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None: pgSz = OxmlElement('w:pgSz'); sectPr.append(pgSz)
    pgSz.set(qn('w:orient'), 'landscape')

    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    table = _new_shelf_table(doc)
    for i, item in enumerate(items):
        integer, decimal = split_price(item["price"])
        _add_shelf_row(table, item["name"].upper(), integer, decimal, currency)
        if (i + 1) % 6 == 0 and i < len(items) - 1:
            page_break(doc); table = _new_shelf_table(doc)

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ── Two items landscape ────────────────────────────────────────────────────
def create_a5_landscape_two(items, currency="د.إ"):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.0); s.page_height = Cm(14.8)
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(1.0)
    sectPr = s._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None: pgSz = OxmlElement('w:pgSz'); sectPr.append(pgSz)
    pgSz.set(qn('w:orient'), 'landscape')

    for i in range(0, len(items), 2):
        pair = items[i:i+2]
        for j, item in enumerate(pair):
            name = item["name"].upper()
            integer, decimal = split_price(item["price"])
            unit = item.get("unit","").upper()

            spacer = doc.add_paragraph()
            pPrS = spacer._p.get_or_add_pPr()
            spS = OxmlElement('w:spacing')
            spS.set(qn('w:before'),'0'); spS.set(qn('w:after'),'0')
            spS.set(qn('w:line'),'720'); spS.set(qn('w:lineRule'),'exact')
            pPrS.append(spS)

            p_price = doc.add_paragraph(); set_para_center_no_space(p_price)
            p_price._p.append(make_run(str(integer),  'Impact', 380, bold=True, color_rgb=(0xEE,0x00,0x00)))
            p_price._p.append(make_run(f'.{decimal}', 'Impact', 100, bold=True, color_rgb=(0xEE,0x00,0x00)))
            if unit:
                p_price._p.append(make_run(f' /{unit}','Impact', 80, bold=True, color_rgb=(0x33,0x33,0x33)))
            p_price._p.append(make_run(f' {currency}','aed',    72, bold=True, color_rgb=(0x00,0x00,0x00)))

            p_name = doc.add_paragraph(); set_para_center_no_space(p_name)
            p_name._p.append(make_run(name, 'Impact', 60, bold=True, color_rgb=(0x00,0x00,0x00)))

            if j == 0 and len(pair) == 2:
                div = doc.add_paragraph()
                pPrD = div._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bot = OxmlElement('w:bottom')
                bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
                bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'AAAAAA')
                pBdr.append(bot); pPrD.append(pBdr)

        if i + 2 < len(items):
            page_break(doc)

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ── Routes ─────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        poster_type = request.form.get('poster_type')
        currency    = request.form.get('currency', 'د.إ')
        items       = []

        # Excel upload
        file = request.files.get('excel_file')
        if file and file.filename:
            df = pd.read_excel(file)
            for _, row in df.iterrows():
                items.append({
                    'name':  str(row.get('Description','')),
                    'price': float(row.get('Price', 0)),
                    'unit':  str(row.get('Unit',''))
                })
        else:
            # Manual entry — items sent as JSON fields name_0, price_0, unit_0 ...
            i = 0
            while True:
                name = request.form.get(f'name_{i}')
                if not name:
                    break
                items.append({
                    'name':  name,
                    'price': float(request.form.get(f'price_{i}', 0)),
                    'unit':  request.form.get(f'unit_{i}', '')
                })
                i += 1

        if not items:
            return jsonify({'error': 'No items provided'}), 400

        if poster_type == 'a4_single':
            buf = create_a4_portrait_single(items, currency)
            fname = 'poster_a4.docx'
        elif poster_type == 'shelf':
            buf = create_a5_landscape_shelf(items, currency)
            fname = 'shelf_labels.docx'
        elif poster_type == 'two_items':
            buf = create_a5_landscape_two(items, currency)
            fname = 'poster_two.docx'
        else:
            return jsonify({'error': 'Unknown poster type'}), 400

        return send_file(
            buf,
            as_attachment=True,
            download_name=fname,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
    
